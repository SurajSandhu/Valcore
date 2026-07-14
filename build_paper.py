"""
Generate the Valcore IEEE conference paper as paper.docx.

Formatting goals:
  - IEEE two-column conference layout (single-column title banner + 2-col body).
  - Times New Roman throughout, IEEE point sizes.
  - Heading 1 / Heading 2 / Normal styles configured to IEEE conventions.
  - Real Word tables (Table Grid), not ASCII.
  - Proper paragraph spacing (compact, single line spacing, first-line indent).
"""
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = os.path.dirname(os.path.abspath(__file__))
FIG = os.path.join(HERE, "figures")

FONT = "Times New Roman"

# ---------------------------------------------------------------- low-level helpers
def set_columns(section, num, space_twips=360):
    """Set the number of newspaper-style columns on a section."""
    sectPr = section._sectPr
    cols = sectPr.find(qn("w:cols"))
    if cols is None:
        cols = OxmlElement("w:cols")
        sectPr.append(cols)
    cols.set(qn("w:num"), str(num))
    cols.set(qn("w:space"), str(space_twips))


def set_margins(section):
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(0.625)
    section.right_margin = Inches(0.625)


def _set_run_font(run, size=None, bold=None, italic=None, smallcaps=False):
    run.font.name = FONT
    # ensure east-asian / cs also map to Times New Roman
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs"):
        rfonts.set(qn(attr), FONT)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold
    if italic is not None:
        run.font.italic = italic
    if smallcaps:
        sc = OxmlElement("w:smallCaps")
        rpr.append(sc)


def para(doc_or_cell, text="", style=None, align=None, size=10, bold=False,
         italic=False, indent=None, space_before=0, space_after=0,
         line=1.0, smallcaps=False):
    p = doc_or_cell.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = line
    if indent is not None:
        pf.first_line_indent = Inches(indent)
    if text:
        r = p.add_run(text)
        _set_run_font(r, size=size, bold=bold, italic=italic, smallcaps=smallcaps)
    return p


def configure_styles(doc):
    # Normal
    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal.font.size = Pt(10)
    normal.paragraph_format.line_spacing = 1.0
    normal.paragraph_format.space_after = Pt(0)
    rpr = normal.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs"):
        rfonts.set(qn(attr), FONT)

    # Heading 1 — IEEE section head: centered, small caps, 10pt
    h1 = doc.styles["Heading 1"]
    h1.font.name = FONT
    h1.font.size = Pt(10)
    h1.font.bold = False
    h1.font.italic = False
    h1.font.color.rgb = RGBColor(0, 0, 0)
    h1.font.small_caps = True
    h1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    h1.paragraph_format.space_before = Pt(12)
    h1.paragraph_format.space_after = Pt(4)
    h1.paragraph_format.line_spacing = 1.0

    # Heading 2 — IEEE subsection: italic, left, 10pt
    h2 = doc.styles["Heading 2"]
    h2.font.name = FONT
    h2.font.size = Pt(10)
    h2.font.bold = False
    h2.font.italic = True
    h2.font.color.rgb = RGBColor(0, 0, 0)
    h2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    h2.paragraph_format.space_before = Pt(6)
    h2.paragraph_format.space_after = Pt(2)
    h2.paragraph_format.line_spacing = 1.0


def h1(doc, text):
    p = doc.add_paragraph(style="Heading 1")
    p.add_run(text)
    return p


def h2(doc, text):
    p = doc.add_paragraph(style="Heading 2")
    p.add_run(text)
    return p


def body(doc, text, indent=0.2, justify=True):
    p = para(doc, text, size=10, indent=indent,
             align=WD_ALIGN_PARAGRAPH.JUSTIFY if justify else None,
             space_after=0, line=1.0)
    return p


def table_caption(doc, text):
    p = para(doc, text, align=WD_ALIGN_PARAGRAPH.CENTER, size=8,
             space_before=8, space_after=2, smallcaps=True)
    return p


def fig_caption(doc, text):
    p = para(doc, text, align=WD_ALIGN_PARAGRAPH.CENTER, size=8,
             space_before=4, space_after=8)
    return p


def add_table(doc, headers, rows, col_widths=None, font_size=8):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = True
    # header
    hdr = t.rows[0].cells
    for i, htext in enumerate(headers):
        hdr[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = hdr[i].paragraphs[0].add_run(htext)
        _set_run_font(r, size=font_size, bold=True)
    # data
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cp = cells[i].paragraphs[0]
            cp.alignment = WD_ALIGN_PARAGRAPH.CENTER if i > 0 else WD_ALIGN_PARAGRAPH.LEFT
            r = cp.add_run(str(val))
            _set_run_font(r, size=font_size)
    # tighten cell paragraph spacing
    for row in t.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(1)
                p.paragraph_format.space_before = Pt(1)
                p.paragraph_format.line_spacing = 1.0
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in t.rows:
                row.cells[i].width = Inches(w)
    return t


def add_figure(doc, path, width_in, caption):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run()
    run.add_picture(path, width=Inches(width_in))
    fig_caption(doc, caption)


# ================================================================ build document
doc = Document()
configure_styles(doc)

# Section 1 (banner): margins + single column
sec1 = doc.sections[0]
set_margins(sec1)
set_columns(sec1, 1)

# ---- Title
title = para(doc, "Valcore: Leakage-Aware Machine Learning for "
                  "Generalizable IoT Botnet Detection",
             align=WD_ALIGN_PARAGRAPH.CENTER, size=24, space_after=6, line=1.0)

# ---- Authors (anonymized for double-blind review)
para(doc, "Anonymous Author(s)", align=WD_ALIGN_PARAGRAPH.CENTER, size=11,
     space_after=0)
para(doc, "Author names and affiliations omitted for double-blind review",
     align=WD_ALIGN_PARAGRAPH.CENTER, size=10, italic=True, space_after=0)
para(doc, "Paper ID: XXXX", align=WD_ALIGN_PARAGRAPH.CENTER, size=10,
     space_after=6)

# ---- switch to two columns for the body
doc.add_section(WD_SECTION.CONTINUOUS)
sec2 = doc.sections[-1]
set_margins(sec2)
set_columns(sec2, 2, space_twips=360)

# ---- Abstract
ab = doc.add_paragraph()
ab.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
ab.paragraph_format.space_after = Pt(6)
ab.paragraph_format.line_spacing = 1.0
r = ab.add_run("Abstract—")
_set_run_font(r, size=9, bold=True, italic=True)
abstract_text = (
    "Botnets are among the most damaging threats facing Internet-of-Things "
    "(IoT) deployments, and machine-learning classifiers trained on public "
    "capture datasets routinely report accuracy close to 100%. Much of that "
    "performance, we argue, is an artifact of data leakage. A model that "
    "memorizes host identifiers such as IP addresses and ports looks excellent "
    "under a random train/test split, but its accuracy collapses on hosts it "
    "has never encountered. We present Valcore, a Random Forest botnet "
    "detector for the UNSW Bot-IoT dataset that is designed to avoid this "
    "trap. We exclude target-leakage and host-identity columns from the "
    "outset and run a four-stage ablation (Models A-D) under two protocols: a "
    "conventional random split and a host-grouped split whose test hosts do "
    "not overlap with the training hosts. The deployed model, Model C, relies "
    "on just ten behavioral flow features. On unseen hosts it reaches 0.991 "
    "accuracy and 0.992 F1, which matches the strongest leak-free model that "
    "avoids identifiers (Model B) and falls within 0.5% of the "
    "identifier-based baseline, all without touching host or port "
    "information. We package Valcore as a Flask web dashboard that classifies "
    "uploaded flow records and returns a per-file threat assessment. On the "
    "same model, the random split reports 0.995 accuracy, 1.000 precision, and "
    "1.000 ROC-AUC; the gap between those figures and the grouped ones "
    "measures the optimism that host-grouped evaluation removes. We report "
    "this optimism, along with other limitations, to keep the results from "
    "being over-read."
)
r = ab.add_run(abstract_text)
_set_run_font(r, size=9, italic=True)

kw = doc.add_paragraph()
kw.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
kw.paragraph_format.space_after = Pt(6)
r = kw.add_run("Index Terms—")
_set_run_font(r, size=9, bold=True, italic=True)
r = kw.add_run("IoT security; botnet detection; intrusion detection; "
               "Random Forest; data leakage; Bot-IoT; generalization.")
_set_run_font(r, size=9, italic=True)

# ================================================================ I. Introduction
h1(doc, "I. Introduction")
body(doc,
     "Low-cost IoT devices ship with weak security and are now everywhere, "
     "which has made them a favored building block for large-scale botnets. "
     "Once compromised, cameras, routers, and sensors are conscripted to "
     "launch distributed denial-of-service attacks, scan for new victims, and "
     "exfiltrate data. The traffic these devices produce is distinctive and "
     "highly repetitive, and that regularity makes supervised machine learning "
     "a natural fit for spotting botnet flows at the network edge.")
body(doc,
     "A recurring problem in this literature is that headline accuracy is "
     "often inflated by data leakage [7]. Public capture datasets come from a "
     "handful of hosts recorded over short windows, so the records within a "
     "single attack burst are near-duplicates of one another. Split them "
     "randomly into train and test sets and almost identical rows land on both "
     "sides. A classifier can then get away with memorizing host "
     "identifiers—IP addresses, ports, or capture-order fields—and post "
     "excellent scores that evaporate once it faces genuinely unseen hosts.")
body(doc,
     "We make three contributions. First, we describe a leakage-aware feature "
     "design that drops target-leakage and host-identity columns from the "
     "start. Second, we run an ablation (Models A-D) that peels away feature "
     "groups one at a time and evaluate each model under both a random and a "
     "host-grouped split, which isolates the real generalization gap. Third, "
     "we wrap the resulting model, Model C, in Valcore—a deployable Flask "
     "dashboard that classifies uploaded flow records and produces an "
     "actionable threat assessment. We also spell out the limitations of our "
     "evaluation, so the numbers are read for what they are.")

# ================================================================ II. Related Work
h1(doc, "II. Related Work")
body(doc,
     "Mirai and its variants showed just how much scale an attacker can reach "
     "by recruiting IoT devices [2], and that demonstration spurred a large "
     "body of learning-based intrusion-detection research. To support it, the "
     "UNSW Bot-IoT dataset was released with realistic labeled IoT attack "
     "traffic and a curated \"10-best-features\" subset aimed at classifier "
     "development [1]. Later work has thrown decision trees, Random Forests, "
     "gradient boosting, and deep neural networks at this dataset and its "
     "relatives, frequently reporting accuracy above 99% [5], [6].")
body(doc,
     "Several authors have warned that these figures can be misleading when "
     "the evaluation ignores host or temporal correlation between the training "
     "and test partitions [7], and they have proposed group-aware and "
     "time-aware validation as fairer alternatives [7]. Valcore sits squarely "
     "in that tradition. We do not propose a new classifier; we take an "
     "ordinary Random Forest and pair it with a leakage-aware feature set and "
     "a host-grouped evaluation, then expose the trained model through a "
     "usable web interface.")

# ================================================================ III. Dataset
h1(doc, "III. Dataset")
body(doc,
     "Our experiments use the UNSW Bot-IoT dataset [1], and specifically the "
     "UNSW_2018_IoT_Botnet_Final_10_Best.csv \"10-best-features\" release—a "
     "semicolon-delimited CSV of labeled IoT network flows. Attack traffic "
     "dominates the raw file. To build a balanced corpus for training and "
     "evaluation, a preprocessing script reads the raw file in fixed-size "
     "chunks, separates normal from attack flows, downsamples the majority "
     "class to a 1:1 ratio, and shuffles with a fixed seed. The result is a "
     "balanced set of 954 records: 477 attack and 477 normal.")
body(doc,
     "A handful of columns are dropped from every model up front, because they "
     "either encode the label outright or pin down specific hosts; Table I "
     "summarizes them. Keeping any of these would let the classifier memorize "
     "identities instead of learning behavior.")

table_caption(doc, "Table I. Columns Excluded From Every Model")
add_table(doc,
          ["Column(s)", "Reason for exclusion"],
          [
              ["pkSeqID", "Row / capture identifier"],
              ["category, subcategory", "Target leakage (encode the label)"],
              ["seq", "Capture-order / temporal label proxy"],
              ["saddr, daddr", "Host-identity leakage (IP addresses)"],
              ["sport, dport", "Environment-specific port identifiers"],
          ],
          col_widths=[1.4, 2.0])

# ================================================================ IV. Methodology
h1(doc, "IV. Methodology")

h2(doc, "A. Leakage-Aware Feature Set")
body(doc,
     "Model C, the deployed model, draws on ten behavioral and statistical "
     "flow features and no host or port identifiers. Dropping those "
     "identifiers is exactly what lets it generalize to hosts that never "
     "appear in training. Table II lists the full feature set. Its one "
     "categorical feature, proto, is label-encoded, and we persist the fitted "
     "encoder so that inference stays reproducible.")

table_caption(doc, "Table II. Model C Feature Set (10 Behavioral Features)")
add_table(doc,
          ["#", "Feature", "Description"],
          [
              ["1", "proto", "Transport/app protocol (label-encoded)"],
              ["2", "stddev", "Std. dev. of aggregated record duration"],
              ["3", "N_IN_Conn_P_SrcIP", "Inbound connections per source IP"],
              ["4", "min", "Minimum duration of aggregated records"],
              ["5", "state_number", "Numerical connection-state encoding"],
              ["6", "mean", "Mean duration of aggregated records"],
              ["7", "N_IN_Conn_P_DstIP", "Inbound connections per dest. IP"],
              ["8", "drate", "Dest.-to-source packets/second rate"],
              ["9", "srate", "Source-to-dest. packets/second rate"],
              ["10", "max", "Maximum duration of aggregated records"],
          ],
          col_widths=[0.25, 1.35, 1.8])

h2(doc, "B. Classifier and Training")
body(doc,
     "The classifier is a Random Forest [3] of 100 trees with a fixed random "
     "seed. We chose it for two reasons: it holds up well on tabular flow "
     "features, and its feature-importance output is easy to interpret. We "
     "split the balanced corpus 80/20, label-encode the categorical proto "
     "feature, and serialize the model, the encoder, and the exact feature "
     "ordering as artifacts that inference later loads. Every model is built "
     "with scikit-learn [4]. Fig. 1 shows the system architecture; Fig. 2 "
     "shows the machine-learning pipeline.")

add_figure(doc, os.path.join(FIG, "fig1_architecture.png"), 3.35,
           "Fig. 1. Valcore system architecture: browser client, Flask "
           "application layer, model artifacts loaded at startup, and the "
           "presentation layer.")
add_figure(doc, os.path.join(FIG, "fig2_ml_pipeline.png"), 3.35,
           "Fig. 2. Machine-learning pipeline from raw Bot-IoT capture through "
           "balancing, leakage-aware feature selection, encoding, Random "
           "Forest training, and persistence/evaluation.")

h2(doc, "C. Inference Contract")
body(doc,
     "At runtime, an uploaded CSV is reindexed onto exactly the Model C "
     "feature set. Extra columns that a raw capture might carry—saddr, daddr, "
     "sport, or dport, for instance—are dropped automatically, and any "
     "unknown categorical value falls back to a known class. Fig. 3 traces "
     "this inference path. Each file comes back summarized as attack and "
     "normal flow counts, a mean model confidence, a discrete threat level "
     "(LOW, MEDIUM, HIGH), and a matching security recommendation.")

add_figure(doc, os.path.join(FIG, "fig3_data_flow.png"), 3.35,
           "Fig. 3. Data-flow diagram of the inference path: upload "
           "validation, preprocessing and encoding, Random Forest "
           "classification, and threat assessment returned to the user.")

# ================================================================ V. Ablation
h1(doc, "V. Ablation Study")
body(doc,
     "To justify the feature set, we run an ablation that starts from a "
     "leak-free baseline (Model A, 14 features) and strips feature groups away "
     "one step at a time. Model B removes the source and destination "
     "addresses; Model C removes the source and destination ports as well; "
     "Model D goes further and drops the two connection-count features. We "
     "evaluate every model under the host-grouped split, where the test hosts "
     "are disjoint from the training hosts. Table III lists accuracy and F1.")

table_caption(doc, "Table III. Ablation Under the Host-Grouped Split")
add_table(doc,
          ["Model", "Features", "Grouped Acc.", "Grouped F1"],
          [
              ["A (incl. IPs)", "14", "0.995", "0.996"],
              ["B (- saddr, daddr)", "12", "0.991", "0.992"],
              ["C (- sport, dport)", "10", "0.991", "0.992"],
              ["D (- conn-count)", "8", "0.967", "0.973"],
          ],
          col_widths=[1.3, 0.7, 0.75, 0.65])
body(doc,
     "Model C holds onto Model B's generalization but with fewer, purely "
     "behavioral features, and it sidesteps Model A's dependence on IP "
     "identifiers. Model D, by contrast, prunes too aggressively and gives up "
     "recall. We therefore deploy Model C.")

# ================================================================ VI. Evaluation
h1(doc, "VI. Evaluation")
body(doc,
     "We put Model C through two protocols. The random 80/20 split is the "
     "convention most prior work reports. The host-grouped split trains and "
     "tests on disjoint sets of source hosts—here five hosts and 213 flow "
     "records that training never touched—and gives the more trustworthy "
     "estimate of generalization. Table IV reports both.")

table_caption(doc, "Table IV. Deployed Model C Performance")
add_table(doc,
          ["Evaluation", "Acc.", "Prec.", "Rec.", "F1", "ROC-AUC"],
          [
              ["Random 80/20 split", "0.995", "1.000", "0.988", "0.994", "1.000"],
              ["Host-grouped (unseen)", "0.991", "0.992", "0.992", "0.992", "1.000"],
          ],
          col_widths=[1.35, 0.55, 0.55, 0.5, 0.5, 0.65])
body(doc,
     "The host-grouped numbers are the ones that matter: on hosts it never saw "
     "in training, Model C reaches 0.991 accuracy and 0.992 F1 from behavioral "
     "features alone. The random-split precision and ROC-AUC of 1.000 are "
     "flattering, propped up by flow-level near-duplicates that straddle the "
     "train and test partitions. Fig. 4 gives the confusion matrix on the "
     "grouped test set, and Fig. 5 gives the Random Forest feature "
     "importances, where the behavioral duration statistics and the "
     "connection-count features carry most of the weight.")

add_figure(doc, os.path.join(FIG, "fig4_confusion_matrix.png"), 2.5,
           "Fig. 4. Model C confusion matrix on the host-grouped test set "
           "(n = 213).")
add_figure(doc, os.path.join(FIG, "fig5_feature_importance.png"), 3.1,
           "Fig. 5. Model C Random Forest Gini feature importances.")

# ================================================================ VII. System
h1(doc, "VII. System Implementation")
body(doc,
     "Valcore is built as a Flask [8] web application. On startup it loads the "
     "serialized model, the label encoder, and the feature-order "
     "specification. From the dashboard a user uploads a comma-delimited CSV "
     "of flow records, and the application validates it—catching wrong "
     "delimiters, empty files, and missing required columns, each with a "
     "specific error message—then preprocesses and encodes the rows, runs "
     "Random Forest inference, and renders a report. That report lists the "
     "number of flows scanned, the attack and normal counts, the mean model "
     "confidence, a threat level, and a recommendation, alongside a doughnut "
     "chart of the attack/normal split. The threat level follows from the "
     "fraction of flows flagged as attacks: under 30% is LOW (keep "
     "monitoring), 30-70% is MEDIUM (investigate suspicious devices), and over "
     "70% is HIGH (isolate the affected devices and start incident "
     "response).")

# ================================================================ VIII. Limitations
h1(doc, "VIII. Limitations")
body(doc,
     "We list several limitations so the results are not over-read. First, the "
     "random-split metrics are optimistic: same-host attack bursts create "
     "near-duplicate flows that leak across a random partition, so the "
     "host-grouped numbers are the ones to trust. Second, the evaluation set "
     "is small and artificially balanced—954 rows in the balanced corpus, only "
     "5 hosts and 213 rows in the grouped test set, and a single seeded split "
     "with no cross-validation or confidence intervals—which leaves the point "
     "estimates fragile. Third, since the holdout is balanced 1:1, the "
     "reported accuracy, and precision above all, will not carry over directly "
     "to a deployment where attacks are rare. Fourth, the balancing step "
     "relies on a 517 MB raw source that we do not ship with the artifact, so "
     "it cannot be reproduced end-to-end from the released corpus alone. "
     "Finally, Valcore is a research and demonstration artifact: it has no "
     "live-capture ingestion, authentication, rate limiting, or model "
     "monitoring, and would need hardening before any operational use.")

# ================================================================ IX. Conclusion
h1(doc, "IX. Conclusion")
body(doc,
     "We presented Valcore, a leakage-aware IoT botnet detector built on the "
     "UNSW Bot-IoT dataset. By leaving out target-leakage and host-identity "
     "features and by testing under a host-grouped split, we showed that a "
     "Random Forest with just ten behavioral flow features generalizes to "
     "unseen hosts at 0.991 accuracy and 0.992 F1—matching the best leak-free "
     "baseline that avoids identifiers (Model B) and landing within 0.5% of "
     "the identifier-based baseline (Model A), all without leaning on IP "
     "addresses or ports. "
     "The model runs behind a practical Flask dashboard that returns an "
     "actionable per-file threat assessment. Future work includes time-aware "
     "and cross-dataset evaluation, calibration under realistic class "
     "prevalence, and streaming inference for live network monitoring.")

# ================================================================ References
h1(doc, "References")
refs = [
    "N. Koroniotis, N. Moustafa, E. Sitnikova, and B. Turnbull, \"Towards the "
    "development of realistic botnet dataset in the Internet of Things for "
    "network forensic analytics: Bot-IoT dataset,\" Future Generation Computer "
    "Systems, vol. 100, pp. 779-796, 2019.",
    "M. Antonakakis et al., \"Understanding the Mirai botnet,\" in Proc. 26th "
    "USENIX Security Symposium, 2017, pp. 1093-1110.",
    "L. Breiman, \"Random forests,\" Machine Learning, vol. 45, no. 1, "
    "pp. 5-32, 2001.",
    "F. Pedregosa et al., \"Scikit-learn: Machine learning in Python,\" "
    "Journal of Machine Learning Research, vol. 12, pp. 2825-2830, 2011.",
    "Y. Meidan et al., \"N-BaIoT: Network-based detection of IoT botnet "
    "attacks using deep autoencoders,\" IEEE Pervasive Computing, vol. 17, "
    "no. 3, pp. 12-22, 2018.",
    "S. Garcia, A. Parmisano, and M. J. Erquiaga, \"IoT-23: A labeled dataset "
    "with malicious and benign IoT network traffic,\" Stratosphere Laboratory, "
    "Tech. Rep., 2020.",
    "A. Hanif Halim et al., \"On the generalization of network intrusion "
    "detection: The impact of evaluation methodology,\" IEEE Access, 2022.",
    "M. Grinberg, Flask Web Development: Developing Web Applications with "
    "Python, 2nd ed. O'Reilly Media, 2018.",
]
for i, ref in enumerate(refs, 1):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.left_indent = Inches(0.2)
    p.paragraph_format.first_line_indent = Inches(-0.2)
    r = p.add_run(f"[{i}] ")
    _set_run_font(r, size=9)
    r = p.add_run(ref)
    _set_run_font(r, size=9)

out = os.path.join(HERE, "paper.docx")
doc.save(out)
print("Saved", out)
